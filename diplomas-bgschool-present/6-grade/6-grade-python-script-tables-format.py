from docx import Document
import pandas as pd
import csv
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import RGBColor

## Conversion only works for Linux as docx2pdf requires Word
#from docx2pdf import convert
from doc2pdf import convert
import os

# Set color
color = (139, 32, 32)  # Red color in RGB

## Convert word table to csv

def docx_table_to_csv(docx_filename, csv_filename):
    # Load the Word document
    doc = Document(docx_filename)
    
    # Assuming the table you want to convert is the first one in the document
    table = doc.tables[0]
    
    # Open a CSV file for writing
    with open(csv_filename, mode='w', newline='') as file:
        writer = csv.writer(file)
        
        # Iterate through the rows in the table
        for row in table.rows:
            # Extract text from each cell in the row
            row_data = [cell.text for cell in row.cells]
            # Write the row to the CSV file
            writer.writerow(row_data)

## Formatting cells and paragraphs

def set_italic_and_color_run(run, color, fontsize):
    """
    Apply formatting to a run.
    """
    run.italic = True
    run.font.color.rgb = RGBColor(*color)
    run.font.size = Pt(fontsize)
    return run
    
def format_value_in_cell(paragraph,key,value):

     # Split the paragraph at the placeholder
     parts = paragraph.text.split(key)
     # Clear the paragraph
     paragraph.clear()
     italic_run = paragraph.add_run(str(value))
     fontsize = 9
     set_italic_and_color_run(italic_run,color,fontsize)


def format_value_in_paragraph(paragraph,key,value):
    # Split the paragraph at the placeholder
    parts = paragraph.text.split(key)
    # Clear the paragraph
    paragraph.clear()
    # Add the text before the placeholder
    if parts[0]:
        paragraph.add_run(parts[0])
    # Add the italicized new text
    italic_run = paragraph.add_run(str(value))
    fontsize = 14
    set_italic_and_color_run(italic_run,color,fontsize)
    # Add the text after the placeholder
    if parts[1]:
        paragraph.add_run(parts[1])


def convert_docx_to_pdf(input_path, output_path=None):
    try:
        # Check if the input file exists
        if not os.path.exists(input_path):
            raise FileNotFoundError(f"The file {input_path} does not exist.")
        
        # Perform the conversion
        if output_path:
            convert(input_path, output_path)
        else:
            convert(input_path)
        
        print(f"Successfully converted {input_path} to PDF.")
    
    except Exception as e:
        print(f"An error occurred: {e}")


## Read each row in data file and replace in the respective paragraph (name, birthdate) 
## OR table (marks, prof names) in the template


def populate_template(row):
    doc = Document(template_path)  # Load the template for each new document
    doc_number = ''
    name = ''
    filename = ''
    for key, value in row.items():
        print('Key: '+key+' Value: ' + str(value))

        # In the text of the certificate
        for paragraph in doc.paragraphs:
            if key in paragraph.text:
                if key == "Document" or key == "Delivery":
                    paragraph.text = paragraph.text.replace(key, str(value))
                else :
                    format_value_in_paragraph(paragraph,key,value)

                # Needed to generate the filename
                if key == 'Document':
                    doc_number = str(value)
                elif key == 'Name':
                    name = str(value)

        filename = doc_number+'-'+name

        # In the tables in the certificate
        for table in doc.tables:
            for cell in table._cells:
                #if f'{{{{{key}}}}}' in cell.text:
                if key in cell.text:					
                    #cell.text = cell.text.replace(key, str(value))
                    for paragraph in cell.paragraphs:
                        if key == 'Bulgarian':
                            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            format_value_in_cell(paragraph,key,value)
                        elif key == 'History':
                            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            format_value_in_cell(paragraph,key,value)
                        elif key == 'Geography':
                            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            format_value_in_cell(paragraph,key,value)
                        else:
                            format_value_in_cell(paragraph,key,value)
                        
                            
    return doc,filename



# Input word and output csv files
docx_filename = '6-grade-data-table-word.docx'  # Replace with the path to your .docx file
csv_filename = docx_filename+'.csv'     # Replace with the desired path for the .csv file

# Covnvert word table to csv
docx_table_to_csv(docx_filename, csv_filename)

# Load the cvs data
data = pd.read_csv(csv_filename)  # Adjust the file path and format as needed

# Load the template document with placeholders
template_path = '6-grade-template.docx'

# Generate documents
for idx, row in data.iterrows():
    (filled_doc,filename) = populate_template(row)
    output_path = filename+'.docx'  # Customize output path and filename
    filled_doc.save(output_path)
    print(f'Saved: {output_path}')
    output_pdf = filename+'.pdf'
    convert_docx_to_pdf(output_path,output_pdf)
    print(f'Saved PDF file: {output_pdf}')

print('All documents generated.')

